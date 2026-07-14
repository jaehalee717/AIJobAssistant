"""
report_generator.py
AIJobAssistant
Version : v1.2.2
"""

import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt

from models.job import Job

from modules.report.summary_writer import SummaryWriter
from modules.report.analysis_writer import AnalysisWriter
from modules.report.mail_writer import MailWriter


class ReportGenerator:

    def __init__(self, report_folder: str):

        self.report_folder = Path(report_folder)
        self.report_folder.mkdir(
            parents=True,
            exist_ok=True,
        )

        self.execution_time = datetime.now()

        self.doc = Document()

        title = self.doc.add_heading(
            "AI Job Assistant Report",
            level=1,
        )

        title.runs[0].font.size = Pt(18)

        self.header = self.doc.add_paragraph()

        self.job_count = 0

    @staticmethod
    def safe_filename(text: str) -> str:

        if not text:
            return "Unknown"

        text = re.sub(r"<[^>]+>", "", text)
        text = re.sub(r'[\\/:*?"<>|]', "_", text)
        text = re.sub(r"\s+", " ", text).strip()

        return text[:60] or "Unknown"

    def generate(self, job: Job):

        self.job_count += 1

        self.doc.add_heading(
            f"Job #{self.job_count}",
            level=2,
        )

        SummaryWriter.write(
            self.doc,
            job,
        )

        AnalysisWriter.write(
            self.doc,
            job,
        )

        MailWriter.write(
            self.doc,
            job,
        )

    def save(
        self,
        processed: int,
        skipped: int,
        failed: int,
    ) -> str:

        self.header.clear()

        self.header.add_run(
            f"Execution Time : "
            f"{self.execution_time:%Y-%m-%d %H:%M:%S}\n"
        )

        self.header.add_run(
            f"Processed : {processed}\n"
        )

        self.header.add_run(
            f"Skipped : {skipped}\n"
        )

        self.header.add_run(
            f"Failed : {failed}"
        )

        filename = (
            "AIJobAssistant_"
            f"{self.execution_time:%Y%m%d_%H%M}"
            ".docx"
        )

        filepath = self.report_folder / filename

        self.doc.save(filepath)

        return str(filepath)