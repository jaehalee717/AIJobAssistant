"""
modules/report/cl_writer.py

AIJobAssistant
Version : v8.1.0
"""

from docx import Document
from docx.shared import Pt

from .cl_parser import CLParser


class CLWriter:

    @classmethod
    def write(
        cls,
        template,
        output,
        ai_text,
    ) -> Document:

        context = CLParser.parse(
            ai_text,
        )

        template_doc = Document(
            template,
        )

        document = Document()

        cls._copy_page_setup(
            template_doc,
            document,
        )

        cls._copy_style(
            template_doc,
            document,
        )

        cls._add_paragraph(
            document,
            context["greeting"],
        )

        for text in context["body"].split("\n"):

            text = text.strip()

            if not text:
                continue

            cls._add_paragraph(
                document,
                text,
            )

        copy_mode = False

        for paragraph in template_doc.paragraphs:

            if paragraph.text.strip() == "Best regards,":

                copy_mode = True

            if copy_mode:

                new_p = document.add_paragraph()

                for run in paragraph.runs:

                    new_run = new_p.add_run(
                        run.text,
                    )

                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline

                    new_run.font.name = (
                        run.font.name
                        or "Calibri"
                    )

                    if run.font.size:

                        new_run.font.size = (
                            run.font.size
                        )

                    else:

                        new_run.font.size = Pt(10.5)

        document.save(
            output,
        )

        return document

    @staticmethod
    def _add_paragraph(
        document,
        text,
    ):

        paragraph = document.add_paragraph()

        run = paragraph.add_run(
            text,
        )

        run.font.name = "Calibri"
        run.font.size = Pt(10.5)

    @staticmethod
    def _copy_page_setup(
        source,
        target,
    ):

        s = source.sections[0]
        t = target.sections[0]

        t.left_margin = s.left_margin
        t.right_margin = s.right_margin
        t.top_margin = s.top_margin
        t.bottom_margin = s.bottom_margin
        t.page_width = s.page_width
        t.page_height = s.page_height

    @staticmethod
    def _copy_style(
        source,
        target,
    ):

        normal = target.styles["Normal"]

        source_normal = source.styles["Normal"]

        normal.font.name = (
            source_normal.font.name
        )

        if source_normal.font.size:

            normal.font.size = (
                source_normal.font.size
            )