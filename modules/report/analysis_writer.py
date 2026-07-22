"""
modules/report/analysis_writer.py

AIJobAssistant
Version : v4.0.0
"""

from __future__ import annotations

from docx.document import Document


class AnalysisWriter:

    @staticmethod
    def write(
        document: Document,
        jobs,
    ) -> None:

        document.add_heading(
            "Analysis",
            level=1,
        )

        for job in jobs:

            document.add_heading(
                f"{job.company} - {job.position}",
                level=2,
            )

            document.add_paragraph(
                f"Company : {job.company}"
            )

            document.add_paragraph(
                f"Position : {job.position}"
            )

            document.add_paragraph(
                f"Location : {job.location}"
            )

            document.add_paragraph(
                f"Salary : {getattr(job, 'salary', 'Unknown')}"
            )

            document.add_paragraph(
                f"Match Score : {getattr(job, 'match', 'Unknown')}"
            )

            document.add_paragraph(
                f"Recommendation : {getattr(job, 'decision', 'Unknown')}"
            )

            document.add_paragraph(
                f"Reason : {getattr(job, 'reason', '')}"
            )

            document.add_paragraph(
                f"Apply URL : {getattr(job, 'apply_url', '')}"
            )

            document.add_paragraph(
                f"Generated Time : {getattr(job, 'created_at', '')}"
            )

            document.add_page_break()