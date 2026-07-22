"""
modules/report/detail_report_writer.py

AIJobAssistant
Version : v4.0.0
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from .detail_writer import DetailWriter
from .summary_writer import SummaryWriter


class DetailReportWriter:
    """Detail Analysis Report 생성"""

    @staticmethod
    def write(
        jobs,
        output_file: Path,
    ) -> Path:

        document = Document()

        document.add_heading(
            "AIJobAssistant Detail Report",
            level=0,
        )

        document.add_paragraph(
            f"Generated Time : {datetime.now():%Y-%m-%d %H:%M:%S}"
        )

        SummaryWriter.write(
            document,
            jobs,
        )

        DetailWriter.write(
            document,
            jobs,
        )

        output_file.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document.save(
            output_file,
        )

        return output_file