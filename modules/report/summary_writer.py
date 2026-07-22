"""
modules/report/summary_writer.py

AIJobAssistant
Version : v4.0.0
"""

from __future__ import annotations

from datetime import datetime

from docx.document import Document


class SummaryWriter:

    @staticmethod
    def write(
        document: Document,
        jobs,
    ) -> None:

        total = len(jobs)
        processed = total
        skipped = sum(
            1 for job in jobs
            if getattr(job, "decision", "") == "SKIP"
        )
        failed = sum(
            1 for job in jobs
            if getattr(job, "status", "") == "FAILED"
        )

        document.add_heading(
            "Summary",
            level=1,
        )

        document.add_paragraph(
            f"Generated Time : {datetime.now():%Y-%m-%d %H:%M:%S}"
        )

        document.add_paragraph(
            "Execution Time : N/A"
        )

        document.add_paragraph(
            f"Processed : {processed}"
        )

        document.add_paragraph(
            f"Skipped : {skipped}"
        )

        document.add_paragraph(
            f"Failed : {failed}"
        )

        document.add_page_break()