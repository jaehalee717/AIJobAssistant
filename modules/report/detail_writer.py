"""
modules/report/detail_writer.py

AIJobAssistant
Version : v4.0.0
"""

from datetime import datetime

from docx import Document


class DetailWriter:
    """Detail Analysis Report 작성"""

    @staticmethod
    def write(
        document: Document,
        jobs,
    ) -> None:

        document.add_heading(
            "Detail Analysis",
            level=1,
        )

        for job in jobs:

            detail = getattr(
                job,
                "detail_result",
                None,
            )

            document.add_heading(
                f"{job.company} | {job.position}",
                level=2,
            )

            document.add_paragraph(
                f"Company : {job.company}"
            )

            document.add_paragraph(
                f"Position : {job.position}"
            )

            document.add_paragraph(
                f"Location : {job.location}"
            )

            document.add_paragraph(
                f"Apply URL : {job.apply_url}"
            )

            document.add_paragraph(
                f"Generated Time : {datetime.now():%Y-%m-%d %H:%M:%S}"
            )

            if detail is None:

                document.add_paragraph(
                    "Detail Analysis : N/A",
                )

                document.add_page_break()

                continue

            fields = [
                ("Recommendation", detail.recommendation),
                ("Priority", detail.priority),
                ("ATS Match", detail.ats_match),
                ("Interview Probability", detail.interview_probability),
                ("Expected Salary", detail.expected_salary),
                ("Employment Type", detail.employment_type),
                ("Work Model", detail.work_model),
                ("Language Fit", detail.language_fit),
                ("Visa / Work Authorization", detail.visa),
                ("Technical Fit", detail.technical_fit),
                ("Management Fit", detail.management_fit),
                ("Leadership Fit", detail.leadership_fit),
                ("Can Perform", detail.can_perform),
                ("Core Strengths", detail.core_strengths),
                ("Technical Gaps", detail.technical_gaps),
                ("Risk", detail.risk),
                ("CV Focus", detail.cv_focus),
                ("Cover Letter Focus", detail.cover_letter_focus),
                ("Reason", detail.reason),
                (
                    "Overall Comments",
                    getattr(
                        detail,
                        "overall_comments",
                        "",
                    ),
                ),
            ]

            for title, value in fields:

                document.add_paragraph(
                    f"{title} : {value}"
                )

            document.add_page_break()