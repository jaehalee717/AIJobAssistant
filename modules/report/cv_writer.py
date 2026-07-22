"""
modules/report/cv_writer.py

AIJobAssistant
Version : v6.2.1
"""

from docx import Document

from .cv_parser import CVParser


class CVWriter:

    @classmethod
    def write(
        cls,
        template,
        output,
        ai_text,
    ) -> None:

        document = Document(
            template,
        )

        context = CVParser.parse(
            ai_text,
        )

        print(context["tai_experience"])
        print("=" * 80)
        print(context["brazil_title"])
        print("=" * 80)

        cls._replace_document(
            document,
            context,
        )

        document.save(
            output,
        )

    @classmethod
    def _replace_document(
        cls,
        document,
        context,
    ) -> None:

        for paragraph in document.paragraphs:

            cls._replace_paragraph(
                paragraph,
                context,
            )

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        cls._replace_paragraph(
                            paragraph,
                            context,
                        )

    @staticmethod
    def _replace_paragraph(
        paragraph,
        context,
    ) -> None:

        key = (
            paragraph.text
            .strip()
            .replace("{{", "")
            .replace("}}", "")
            .strip()
            .lower()
        )

        if key not in context:
            return

        value = context[key]

        if not paragraph.runs:

            run = paragraph.add_run(
                value,
            )

            if key == "professional_profile":
                run.bold = False

            return

        paragraph.runs[0].text = value

        if key == "professional_profile":
            paragraph.runs[0].bold = False

        for run in paragraph.runs[1:]:

            run.text = ""