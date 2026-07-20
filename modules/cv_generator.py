"""
modules/cv_generator.py
AIJobAssistant
Version : v2.0.0
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from modules.output_manager import OutputManager


class CVGenerator:
    """CV Generator"""

    def __init__(
        self,
        template_path: Path,
    ):

        self.template_path = Path(
            template_path,
        )

    def generate(
        self,
        output: OutputManager,
        profile: str,
        competencies: str,
        tai: str,
        brazil: str,
        spain: str,
    ) -> Path:
        """
        Generate CV.

        Returns
            Saved docx path
        """

        document = Document(
            self.template_path,
        )

        self._replace_after_heading(
            document,
            "PROFESSIONAL PROFILE",
            profile,
        )

        self._replace_after_heading(
            document,
            "CORE COMPETENCIES",
            competencies,
        )

        self._replace_after_text(
            document,
            "TAI Escuela Universitaria de Artes",
            tai,
        )

        self._replace_after_text(
            document,
            "LG Electronics Brazil Ltd",
            brazil,
        )

        self._replace_after_text(
            document,
            "LG Electronics Spain SA",
            spain,
        )

        output_path = output.get_cv_docx_path()

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document.save(
            output_path,
        )

        return output_path

    def _replace_after_heading(
        self,
        document: Document,
        heading: str,
        text: str,
    ) -> None:

        paragraphs = document.paragraphs

        for i, p in enumerate(
            paragraphs,
        ):

            if p.text.strip().upper() != heading.upper():
                continue

            if i + 1 >= len(paragraphs):
                return

            paragraphs[i + 1].text = text
            return

    def _replace_after_text(
        self,
        document: Document,
        keyword: str,
        text: str,
    ) -> None:

        paragraphs = document.paragraphs

        for i, p in enumerate(
            paragraphs,
        ):

            if keyword.lower() not in p.text.lower():
                continue

            j = i + 1

            while j < len(paragraphs):

                if paragraphs[j].text.strip().startswith("•"):
                    paragraphs[j].text = text
                    return

                j += 1