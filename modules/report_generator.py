"""
report_generator.py
AIJobAssistant
Version : v1.2.0
"""

import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt

from models.job import Job
from utils.word_helper import add_hyperlink

class ReportGenerator:

    def __init__(self, report_folder: str):

        self.report_folder = Path(report_folder)
        self.report_folder.mkdir(
            parents=True,
            exist_ok=True,
        )

        self.execution_time = datetime.now()

        self.doc = Document()

        title = self.doc.add_heading(
            "AI Job Assistant Report",
            level=1,
        )

        title.runs[0].font.size = Pt(18)

        self.header = self.doc.add_paragraph()

        self.job_count = 0

    @staticmethod
    def safe_filename(text: str) -> str:

        if not text:
            return "Unknown"

        text = re.sub(r"<[^>]+>", "", text)
        text = re.sub(r'[\\/:*?"<>|]', "_", text)
        text = re.sub(r"\s+", " ", text).strip()

        if len(text) > 60:
            text = text[:60]

        return text or "Unknown"

    def generate(self, job: Job):

        self.job_count += 1

        self.doc.add_heading(
            f"Job #{self.job_count}",
            level=2,
        )

        table = self.doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        def add_row(name, value):

            cells = table.add_row().cells
            cells[0].text = name
            cells[1].text = "" if value is None else str(value)

        add_row("Company", job.company)
        add_row("Position", job.position)
        add_row("Location", job.location)

        cells = table.add_row().cells
        cells[0].text = "Apply URL"

        add_hyperlink(
            cells[1].paragraphs[0],
            "Open Job Posting",
            job.apply_url,
        )
        
        add_row("Mail Type", job.mail_type)

        self.doc.add_paragraph()

        self.doc.add_heading(
            "Evaluation",
            level=3,
        )

        table = self.doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        add = lambda n, v: (
            (lambda c: (
                setattr(c[0], "text", n),
                setattr(c[1], "text", "" if v is None else str(v))
            ))(table.add_row().cells)
        )

        add("Match", job.match)
        add("Decision", job.decision)
        add("Confidence", job.confidence)
        add("Strength", job.strength)
        add("Weak", job.weak)
        add("Reason", job.reason)

        self.doc.add_paragraph()

        self.doc.add_heading(
            "Mail",
            level=3,
        )

        body = job.description or job.body or ""

        if len(body) > 8000:
            body = body[:8000]

        self.doc.add_paragraph(body)

        self.doc.add_page_break()

    def save(
        self,
        processed: int,
        skipped: int,
        failed: int,
    ) -> str:

        self.header.clear()

        self.header.add_run(
            f"Execution Time : "
            f"{self.execution_time:%Y-%m-%d %H:%M:%S}\n"
        )

        self.header.add_run(
            f"Processed : {processed}\n"
        )

        self.header.add_run(
            f"Skipped : {skipped}\n"
        )

        self.header.add_run(
            f"Failed : {failed}"
        )

        filename = (
            "AIJobAssistant_"
            f"{self.execution_time:%Y%m%d_%H%M}"
            ".docx"
        )

        filepath = self.report_folder / filename

        self.doc.save(filepath)

        return str(filepath)