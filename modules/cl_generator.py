"""
modules/cl_generator.py
AIJobAssistant
Version : v2.0.0
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from modules.output_manager import OutputManager


class CLGenerator:
    """Cover Letter Generator"""

    def __init__(
        self,
        template_path: Path,
    ):

        self.template_path = Path(
            template_path,
        )

    def generate(
        self,
        output: OutputManager,
        letter: str,
    ) -> Path:
        """
        Generate Cover Letter.

        Returns
            Saved docx path
        """

        document = Document(
            self.template_path,
        )

        self._replace_content(
            document,
            letter,
        )

        output_path = output.get_cl_docx_path()

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document.save(
            output_path,
        )

        return output_path

    def _replace_content(
        self,
        document: Document,
        text: str,
    ) -> None:

        paragraphs = document.paragraphs

        if not paragraphs:

            document.add_paragraph(
                text,
            )

            return

        paragraphs[0].text = text

        for paragraph in paragraphs[1:]:

            paragraph.text = ""